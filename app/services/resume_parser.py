"""
PrepVista AI — Resume Parser Service
PDF extraction, sanitization, and structured parsing via LLM.
"""

import io
import re
import structlog
from pypdf import PdfReader
from fastapi import HTTPException

from app.config import get_settings
from app.services.llm import call_llm_json
from app.services.prompts import build_resume_extraction_prompt

logger = structlog.get_logger("prepvista.resume")

# Fix 6 — accepted resume formats beyond PDF. Images and image-only PDFs go
# through OCR; .docx via python-docx; legacy .doc via LibreOffice headless.
_IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif")
SUPPORTED_RESUME_EXTENSIONS = (".pdf", ".docx", ".doc") + _IMAGE_EXTENSIONS

# Quality gate: below this many extracted characters there is not enough signal
# to build a meaningful interview (covers blank scans, failed OCR, near-empty docs).
_MIN_RESUME_TEXT_CHARS = 200

# Max pages to OCR from an image-only PDF — bounds Tesseract cost on huge scans.
_MAX_OCR_PDF_PAGES = 10

_CONTENT_TYPE_EXTENSIONS = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/msword": ".doc",
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/webp": ".webp",
    "image/bmp": ".bmp",
    "image/tiff": ".tiff",
}


def _resume_extension(filename: str, content_type: str | None = None) -> str:
    """Resolve a supported resume extension from filename, then content-type."""
    name = (filename or "").lower().strip()
    for ext in SUPPORTED_RESUME_EXTENSIONS:
        if name.endswith(ext):
            return ext
    ct = (content_type or "").split(";")[0].strip().lower()
    return _CONTENT_TYPE_EXTENSIONS.get(ct, "")

FIELD_BUCKETS = (
    "software_backend_frontend",
    "ai_ml_data",
    "data_science_analytics",
    "electronics_embedded",
    "electrical_core",
    "mechanical_core",
    "civil_core",
    "cybersecurity",
    "business_analyst_operations",
    "design_creative",
    "general_fresher_mixed",
    "non_technical_general",
)

FIELD_KEYWORDS = {
    "ai_ml_data": {
        "ai", "artificial intelligence", "machine learning", "ml", "nlp", "rag", "llm",
        "prompt", "embedding", "retrieval", "ranking", "evaluation", "grounding",
        "hallucination", "classification", "data science", "computer vision",
        # Modern AI/ML frameworks & tooling — strong, unambiguous AI signals that
        # were previously unrecognised, causing AI/RAG resumes to be misclassified
        # as generic software when they also listed Python/FastAPI/an API.
        "pytorch", "tensorflow", "keras", "hugging face", "huggingface", "transformer",
        "transformers", "openai", "groq", "anthropic", "langchain", "llamaindex",
        "fine-tuning", "fine tuning", "generative",
    },
    "software_backend_frontend": {
        "python", "java", "c++", "javascript", "typescript", "react", "next.js",
        "nextjs", "node", "fastapi", "django", "flask", "spring", "backend",
        "frontend", "api", "database", "sql", "postgresql", "mysql", "mongodb",
        "full stack", "fullstack", "web", "software engineer", "developer",
    },
    "electronics_embedded": {
        "embedded", "microcontroller", "arduino", "raspberry pi", "raspberry",
        "pcb", "firmware", "sensor", "uart", "spi", "i2c", "signal processing",
        "communication systems", "vlsi", "fpga", "verilog", "electronics",
    },
    "electrical_core": {
        "electrical", "power systems", "power electronics", "machines",
        "renewable energy", "control systems", "circuits", "distribution",
        "generation", "transformer", "motor", "grid",
    },
    "mechanical_core": {
        "mechanical", "cad", "solidworks", "autocad", "manufacturing",
        "thermodynamics", "production", "machine design", "automotive",
        "ansys", "cnc", "robotics", "maintenance",
    },
    "civil_core": {
        "civil", "construction", "structural", "site", "autocad civil",
        "surveying", "estimation", "quantity", "bim", "revit", "concrete",
        "transportation engineering",
    },
    "business_analyst_operations": {
        "business analyst", "operations", "process improvement", "excel", "power bi",
        "tableau", "dashboard", "stakeholder", "reporting", "analysis", "analytics",
        "operations analyst", "consulting", "kpi", "process", "finance", "sales ops",
    },
    "design_creative": {
        "design", "ui", "ux", "figma", "prototype", "wireframe", "branding",
        "illustration", "graphic design", "visual design", "product design",
        "user research", "interaction design", "creative",
    },
    "non_technical_general": {
        "marketing", "content", "hr", "human resources", "recruitment", "customer support",
        "teaching", "education", "administration", "operations executive", "communication",
        "public speaking", "sales", "management", "coordination",
    },
    "data_science_analytics": {
        "data science", "data engineering", "data pipeline", "etl", "pandas", "numpy",
        "statistics", "statistical", "visualization", "tableau", "power bi", "bigquery",
        "spark", "hadoop", "data warehouse", "feature engineering", "exploratory data analysis",
        "eda", "jupyter", "matplotlib", "seaborn", "scikit-learn", "regression",
        "hypothesis testing", "a/b testing", "data analyst",
    },
    "cybersecurity": {
        "cybersecurity", "cyber security", "penetration testing", "ethical hacking",
        "vulnerability", "threat", "firewall", "siem", "soc", "incident response",
        "malware", "cryptography", "network security", "information security", "infosec",
        "owasp", "ceh", "security audit", "intrusion detection", "ids", "ips",
        "security operations", "threat modeling", "forensics", "nmap", "wireshark",
        "burp suite", "metasploit",
    },
}

ROLE_LABELS = {
    "ai_ml_data": "AI or ML roles",
    "data_science_analytics": "data science or analytics roles",
    "software_backend_frontend": "software engineering roles",
    "electronics_embedded": "embedded or electronics roles",
    "electrical_core": "electrical engineering roles",
    "mechanical_core": "mechanical engineering roles",
    "civil_core": "civil engineering roles",
    "cybersecurity": "cybersecurity or information security roles",
    "business_analyst_operations": "business analyst or operations roles",
    "design_creative": "design roles",
    "general_fresher_mixed": "entry-level roles aligned with your strongest projects",
    "non_technical_general": "entry-level non-technical or business-facing roles",
}

# ── Prompt injection patterns to strip ───────────────
INJECTION_PATTERNS = [
    r"ignore\s+(all\s+)?previous\s+instructions",
    r"you\s+are\s+now",
    r"system\s*:",
    r"pretend\s+you",
    r"disregard\s+(all\s+)?prior",
    r"new\s+instructions",
    r"forget\s+(everything|all)",
    r"override\s+(your|all)",
]


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract raw text from PDF bytes with pypdf primary, pdfplumber fallback."""
    settings = get_settings()

    # Primary: pypdf (fast, handles most PDFs)
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text.strip())
        raw_text = "\n".join(p for p in pages if p).strip()
        if raw_text and len(raw_text) > 50:
            return raw_text[:settings.MAX_RESUME_TEXT_LENGTH]
    except HTTPException:
        raise
    except Exception as e:
        logger.warning("pypdf_extraction_failed_trying_pdfplumber", error=str(e))

    # Fallback: pdfplumber (handles multi-column, tables, complex layouts)
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages.append(text.strip())
            raw_text = "\n".join(p for p in pages if p).strip()
            if raw_text:
                return raw_text[:settings.MAX_RESUME_TEXT_LENGTH]
    except Exception as e:
        logger.error("pdfplumber_extraction_also_failed", error=str(e))

    # Fix 6 — last resort: OCR an image-only / scanned PDF (best-effort; only
    # runs if pdf2image + pytesseract + their system binaries are available).
    ocr_text = _ocr_pdf(pdf_bytes)
    if ocr_text:
        return ocr_text[:settings.MAX_RESUME_TEXT_LENGTH]

    raise HTTPException(
        status_code=400,
        detail="Could not extract text from the resume PDF. Please ensure it's not an image-only scan, or upload a clearer copy.",
    )


def _ocr_pdf(pdf_bytes: bytes) -> str:
    """OCR a scanned/image-only PDF. Returns "" if OCR is unavailable or fails."""
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError:
        logger.info("pdf_ocr_unavailable", reason="pytesseract/pdf2image not installed")
        return ""
    try:
        images = convert_from_bytes(pdf_bytes, dpi=200, fmt="png")
    except Exception as e:  # poppler missing or corrupt PDF
        logger.warning("pdf_ocr_convert_failed", error=str(e))
        return ""
    chunks: list[str] = []
    for image in images[:_MAX_OCR_PDF_PAGES]:
        try:
            chunks.append(pytesseract.image_to_string(image) or "")
        except Exception as e:  # tesseract binary missing
            logger.warning("pdf_ocr_page_failed", error=str(e))
            break
    return "\n".join(chunk.strip() for chunk in chunks if chunk.strip()).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx resume (paragraphs + tables) via python-docx."""
    settings = get_settings()
    try:
        import docx  # python-docx
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="DOCX resumes are not supported on this server. Please upload a PDF.",
        )
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.warning("docx_extraction_failed", error=str(e))
        raise HTTPException(
            status_code=400,
            detail="Could not read the DOCX resume. Please re-save it as PDF and try again.",
        )
    parts = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()[:settings.MAX_RESUME_TEXT_LENGTH]


def extract_text_from_doc(file_bytes: bytes) -> str:
    """Extract text from a legacy .doc resume via LibreOffice headless conversion."""
    import os
    import shutil
    import subprocess
    import tempfile

    settings = get_settings()
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise HTTPException(
            status_code=503,
            detail="Legacy .doc resumes are not supported on this server. Please upload a PDF or DOCX.",
        )
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "resume.doc")
        with open(src, "wb") as fh:
            fh.write(file_bytes)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", tmpdir, src],
                check=True,
                capture_output=True,
                timeout=60,
            )
        except Exception as e:
            logger.warning("doc_libreoffice_failed", error=str(e))
            raise HTTPException(
                status_code=400,
                detail="Could not convert the .doc resume. Please upload a PDF or DOCX.",
            )
        out = os.path.join(tmpdir, "resume.txt")
        if not os.path.exists(out):
            raise HTTPException(
                status_code=400,
                detail="Could not convert the .doc resume. Please upload a PDF or DOCX.",
            )
        with open(out, "r", encoding="utf-8", errors="ignore") as fh:
            text = fh.read().strip()
    return text[:settings.MAX_RESUME_TEXT_LENGTH]


def extract_text_from_image(file_bytes: bytes) -> str:
    """OCR a resume uploaded as an image (PNG/JPG/etc.) via pytesseract."""
    settings = get_settings()
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="Image resumes are not supported on this server. Please upload a PDF.",
        )
    try:
        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image) or ""
    except Exception as e:
        logger.warning("image_ocr_failed", error=str(e))
        raise HTTPException(
            status_code=400,
            detail="Could not read text from the image. Please upload a clearer scan or a PDF.",
        )
    return text.strip()[:settings.MAX_RESUME_TEXT_LENGTH]


def extract_text_from_resume(file_bytes: bytes, filename: str, content_type: str | None = None) -> str:
    """Fix 6 — format-aware resume text extraction with a <200-char quality gate.

    Routes PDF/DOCX/DOC/image uploads to the right extractor (image-only PDFs and
    image uploads go through OCR) and rejects anything that yields too little
    text to build a meaningful interview.
    """
    ext = _resume_extension(filename, content_type)
    if ext == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        text = extract_text_from_docx(file_bytes)
    elif ext == ".doc":
        text = extract_text_from_doc(file_bytes)
    elif ext in _IMAGE_EXTENSIONS:
        text = extract_text_from_image(file_bytes)
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported resume format. Please upload one of: {', '.join(SUPPORTED_RESUME_EXTENSIONS)}.",
        )

    text = (text or "").strip()
    if len(text) < _MIN_RESUME_TEXT_CHARS:
        raise HTTPException(
            status_code=400,
            detail=(
                "The extracted resume text was too short to build a meaningful interview "
                f"(under {_MIN_RESUME_TEXT_CHARS} characters). Please upload a more complete "
                "resume or a clearer scan."
            ),
        )
    return text


def validate_resume_upload(file_bytes: bytes, filename: str, content_type: str | None = None) -> str:
    """Validate a multi-format resume upload (size + extension + content magic).

    Returns the resolved extension. Mirrors validate_pdf_upload but for all
    formats accepted by extract_text_from_resume.
    """
    settings = get_settings()
    if len(file_bytes) > settings.MAX_RESUME_SIZE_BYTES:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is {settings.MAX_RESUME_SIZE_BYTES // (1024*1024)}MB.",
        )
    ext = _resume_extension(filename, content_type)
    if ext not in SUPPORTED_RESUME_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file extension. Please upload one of: {', '.join(SUPPORTED_RESUME_EXTENSIONS)}.",
        )
    # Content-level magic-byte checks defeat extension spoofing.
    if ext == ".pdf" and not file_bytes.startswith(b"%PDF"):
        raise HTTPException(status_code=400, detail="Invalid file. The .pdf does not look like a real PDF.")
    if ext == ".docx" and not file_bytes.startswith(b"PK\x03\x04"):  # OOXML is a zip
        raise HTTPException(status_code=400, detail="Invalid file. The .docx does not look like a real Word document.")
    if ext == ".doc" and not file_bytes.startswith(b"\xd0\xcf\x11\xe0"):  # OLE2 compound file
        raise HTTPException(status_code=400, detail="Invalid file. The .doc does not look like a real Word document.")
    if ext in _IMAGE_EXTENSIONS:
        image_magics = (b"\x89PNG", b"\xff\xd8\xff", b"RIFF", b"BM", b"II*\x00", b"MM\x00*")
        if not any(file_bytes.startswith(magic) for magic in image_magics):
            raise HTTPException(status_code=400, detail="Invalid file. The image does not look like a real image.")
    return ext


def sanitize_resume_text(text: str) -> str:
    """Strip potential prompt injection patterns from resume text."""
    sanitized = text
    for pattern in INJECTION_PATTERNS:
        sanitized = re.sub(pattern, "[filtered]", sanitized, flags=re.IGNORECASE)
    return sanitized


def validate_pdf_upload(file_bytes: bytes, filename: str):
    """Validate the uploaded PDF file."""
    settings = get_settings()

    if len(file_bytes) > settings.MAX_RESUME_SIZE_BYTES:
        raise HTTPException(status_code=400, detail=f"File too large. Maximum size is {settings.MAX_RESUME_SIZE_BYTES // (1024*1024)}MB.")

    if not file_bytes.startswith(b"%PDF"):
        raise HTTPException(status_code=400, detail="Invalid file type. Must be a valid PDF.")

    if not filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Invalid file extension. Must be a .pdf file.")


async def parse_resume_structured(resume_text: str) -> dict:
    """Extract structured resume data using LLM."""
    try:
        prompt = build_resume_extraction_prompt(resume_text)
        result = await call_llm_json(
            [{"role": "system", "content": prompt}],
            temperature=0.1,
        )
        # Validate minimum structure
        if not isinstance(result, dict):
            return _default_resume_summary(resume_text)
        return enrich_resume_summary(result, resume_text=resume_text)
    except Exception as e:
        logger.warning("resume_extraction_failed", error=str(e))
        return _default_resume_summary(resume_text)


def _collect_resume_signal_parts(summary: dict) -> tuple[str, list[str]]:
    """Flatten structured resume sections for deterministic field inference."""
    pieces: list[str] = []
    signal_sources: list[str] = []

    inferred_role = str(summary.get("inferred_role") or "").strip()
    if inferred_role:
        pieces.append(inferred_role)
        signal_sources.append("inferred_role")

    for key in ("skills", "education"):
        values = summary.get(key) or []
        if isinstance(values, list):
            cleaned_values = [str(value).strip() for value in values if str(value).strip()]
            if cleaned_values:
                pieces.extend(cleaned_values)
                signal_sources.append(key)

    for key in ("projects", "experience"):
        values = summary.get(key) or []
        if not isinstance(values, list):
            continue
        for item in values:
            if not isinstance(item, dict):
                continue
            for field in ("name", "title", "role", "description"):
                value = str(item.get(field) or "").strip()
                if value:
                    pieces.append(value)
                    signal_sources.append(key)
            tech_stack = item.get("tech_stack") or []
            if isinstance(tech_stack, list):
                cleaned_stack = [str(value).strip() for value in tech_stack if str(value).strip()]
                if cleaned_stack:
                    pieces.extend(cleaned_stack)
                    signal_sources.append(key)

    return " ".join(pieces).lower(), list(dict.fromkeys(signal_sources))


def infer_resume_field_profile(summary: dict | None) -> dict:
    """Infer a broad resume field and lightweight confidence using deterministic signals."""
    summary = summary if isinstance(summary, dict) else {}
    resume_blob, signal_sources = _collect_resume_signal_parts(summary)
    if not resume_blob:
        return {
            "broad_field": "general_fresher_mixed",
            "field_confidence": 0.2,
            "target_role_label": ROLE_LABELS["general_fresher_mixed"],
            "strong_signal_sources": signal_sources,
        }

    scores = {bucket: 0 for bucket in FIELD_BUCKETS}
    for bucket, keywords in FIELD_KEYWORDS.items():
        for keyword in keywords:
            if keyword in resume_blob:
                scores[bucket] += 1

    education_text = " ".join(str(item).lower() for item in (summary.get("education") or []) if isinstance(item, str))
    if "computer science" in education_text or "information technology" in education_text:
        scores["software_backend_frontend"] += 1
    if "artificial intelligence" in education_text or "data science" in education_text:
        scores["ai_ml_data"] += 1
    if "electronics" in education_text or "embedded" in education_text:
        scores["electronics_embedded"] += 1
    if "electrical" in education_text:
        scores["electrical_core"] += 1
    if "mechanical" in education_text:
        scores["mechanical_core"] += 1
    if "civil" in education_text:
        scores["civil_core"] += 1
    if any(term in education_text for term in ["business", "commerce", "management", "operations"]):
        scores["business_analyst_operations"] += 1
    if any(term in education_text for term in ["design", "visual communication", "fine arts"]):
        scores["design_creative"] += 1
    if any(term in education_text for term in ["data science", "data analytics", "statistics", "data engineering"]):
        scores["data_science_analytics"] += 1
    if any(term in education_text for term in ["cybersecurity", "cyber security", "information security", "network security"]):
        scores["cybersecurity"] += 1

    ranked = sorted(scores.items(), key=lambda item: item[1], reverse=True)
    top_bucket, top_score = ranked[0]
    second_score = ranked[1][1] if len(ranked) > 1 else 0

    if top_score <= 1:
        top_bucket = "general_fresher_mixed"
    elif top_score <= 2 and second_score > 0 and (top_score - second_score) <= 1:
        top_bucket = "general_fresher_mixed"

    confidence = round(min(1.0, max(0.2, top_score / max(top_score + second_score, 2))), 2)
    return {
        "broad_field": top_bucket,
        "field_confidence": confidence,
        "target_role_label": ROLE_LABELS.get(top_bucket, ROLE_LABELS["general_fresher_mixed"]),
        "strong_signal_sources": signal_sources[:4],
    }


def enrich_resume_summary(summary: dict | None, resume_text: str = "") -> dict:
    """Preserve current structure while adding deterministic field signals."""
    base = dict(summary or {})
    candidate_name = str(base.get("candidate_name") or "").strip()
    if not candidate_name:
        first_line = resume_text.splitlines()[0].strip() if resume_text.splitlines() else "Unknown"
        base["candidate_name"] = first_line[:80]

    base.setdefault("education", [])
    base.setdefault("skills", [])
    base.setdefault("projects", [])
    base.setdefault("experience", [])
    base.setdefault("inferred_role", "other")
    # ✅ ADDED: first-class fields for the four new question families
    # (programming_language, skill_verification, certification, self_assessment).
    # Previously languages were buried inside `skills` and certifications were not
    # extracted at all, so the planner had no structured signal to probe them.
    base.setdefault("certifications", [])
    base.setdefault("programming_languages", [])

    field_profile = infer_resume_field_profile(base)
    for key, value in field_profile.items():
        base[key] = value

    return base


def _default_resume_summary(resume_text: str) -> dict:
    """Fallback resume summary when LLM extraction fails."""
    first_line = resume_text.splitlines()[0].strip() if resume_text.splitlines() else "Unknown"
    return enrich_resume_summary({
        "candidate_name": first_line[:80],
        "education": [],
        "skills": [],
        "projects": [],
        "experience": [],
        "inferred_role": "other",
        "certifications": [],
        "programming_languages": [],
    }, resume_text=resume_text)
